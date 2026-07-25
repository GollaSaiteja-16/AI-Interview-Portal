from fastapi import UploadFile, File
import fitz
import speech_recognition as sr
import tempfile
from fastapi.middleware.cors import CORSMiddleware
from fastapi import FastAPI, Depends, UploadFile, File
from sqlalchemy.orm import Session
from database import SessionLocal, engine, Base
from models import User
from models import User, InterviewHistory
from fastapi import FastAPI, Depends, UploadFile, File
from sqlalchemy.orm import Session
from database import SessionLocal, engine, Base
from models import User
import google.generativeai as genai
import os
from dotenv import load_dotenv
from schemas import (
    UserCreate,
    AdminLogin,
    UserLogin,
    ForgotPassword,
    QuestionRequest,
    AnswerRequest,
    AIAnswerRequest,
    HRAnswerRequest,
    CodeAnswerRequest
)

from docx import Document
import fitz
import re
import speech_recognition as sr
import tempfile

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
load_dotenv()

genai.configure(
    api_key=os.getenv("GEMINI_API_KEY")
)

model = genai.GenerativeModel("gemini-3.5-flash")
Base.metadata.create_all(bind=engine)


# Database Connection
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# Home API
@app.get("/")
def home():
    return {"message": "AI Interview Portal API Running"}


# Register API
@app.post("/register")
def register(user: UserCreate, db: Session = Depends(get_db)):

    new_user = User(
        name=user.name,
        email=user.email,
        password=user.password
    )

    db.add(new_user)
    db.commit()

    return {"message": "User Registered Successfully"}


# Admin Login API
@app.post("/admin-login")
def admin_login(admin: AdminLogin):

    if admin.email == "admin@gmail.com" and admin.password == "admin123":
        return {"message": "Admin Login Successful"}

    return {"message": "Invalid Admin Credentials"}


# User Login API
@app.post("/login")
def login(user: UserLogin, db: Session = Depends(get_db)):

    existing_user = db.query(User).filter(
        User.email == user.email,
        User.password == user.password
    ).first()

    if existing_user:
        return {"message": "Login Successful"}

    return {"message": "Invalid Email or Password"}


# Forgot Password API
@app.post("/forgot-password")
def forgot_password(user: ForgotPassword, db: Session = Depends(get_db)):

    existing_user = db.query(User).filter(
        User.email == user.email
    ).first()

    if existing_user:
        return {"message": "Email found. Password reset link sent."}

    return {"message": "Email not registered"}


# Resume Analyzer API
@app.post("/analyze-resume")
async def analyze_resume(file: UploadFile = File(...)):

    text = ""
    name = ""
    email = ""
    phone = ""
    education = ""
    experience = ""
    projects = []

    # PDF Resume
    if file.filename.endswith(".pdf"):

        pdf_data = await file.read()
        pdf = fitz.open(stream=pdf_data, filetype="pdf")

        for page in pdf:
            text += page.get_text()

        lines = text.split("\n")

        if len(lines) > 0:
            name = lines[0].strip()

    # DOCX Resume
    elif file.filename.endswith(".docx"):

        doc = Document(file.file)

        for para in doc.paragraphs:
            text += para.text + "\n"

        lines = text.split("\n")

        if len(lines) > 0:
            name = lines[0].strip()

    # Email Extraction
    email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"

    emails = re.findall(email_pattern, text)

    if emails:
        email = emails[0]

    # Phone Extraction
    phone_pattern = r"\b\d{10}\b"

    phones = re.findall(phone_pattern, text)

    if phones:
        phone = phones[0]

    # Education Extraction
    education_keywords = [
        "Bachelor",
        "B.Tech",
        "BTech",
        "Engineering",
        "Degree",
        "Diploma",
        "Polytechnic",
        "Master",
        "M.Tech"
    ]

    for keyword in education_keywords:
        if keyword.lower() in text.lower():
            education = keyword
            break
        # Experience Extraction

         # Experience Extraction
    experience_keywords = [
        "Internship",
        "Intern",
        "Experience",
        "Worked",
        "Software Developer",
        "Trainee"
    ]

    for keyword in experience_keywords:
        if keyword.lower() in text.lower():
            experience = keyword
            break
@app.post("/generate-questions")
def generate_questions(data: QuestionRequest):

    questions = []

    for skill in data.skills:

        if skill.lower() == "python":
            questions.extend([
                "What is Python?",
                "Explain Python decorators.",
                "Difference between List and Tuple?"
            ])

        elif skill.lower() == "java":
            questions.extend([
                "What is JVM?",
                "Explain OOP concepts in Java.",
                "Difference between JDK and JRE?"
            ])

        elif skill.lower() == "machine learning":
            questions.extend([
                "What is Machine Learning?",
                "What is Overfitting?",
                "Explain Gradient Boosting."
            ])

        elif skill.lower() == "sql":
            questions.extend([
                "What is SQL?",
                "Difference between DELETE and TRUNCATE?",
                "Explain JOINS."
            ])

        elif skill.lower() == "html":
            questions.extend([
                "What is HTML?",
                "Difference between div and span?",
                "What are semantic tags?"
            ])

        elif skill.lower() == "css":
            questions.extend([
                "What is CSS?",
                "Difference between Flexbox and Grid?",
                "Explain Position Property."
            ])

        elif skill.lower() == "javascript":
            questions.extend([
                "What is JavaScript?",
                "Difference between var, let and const?",
                "Explain closures."
            ])

    return {
        "questions": questions
    }
@app.post("/resume-interview")
async def resume_interview(file: UploadFile = File(...)):

    text = ""

    if file.filename.endswith(".pdf"):
        pdf_data = await file.read()
        pdf = fitz.open(stream=pdf_data, filetype="pdf")

        for page in pdf:
            text += page.get_text()

    elif file.filename.endswith(".docx"):
        doc = Document(file.file)

        for para in doc.paragraphs:
            text += para.text + "\n"

    skills = []

    skill_list = [
        "Python",
        "Java",
        "React",
        "Machine Learning",
        "SQL",
        "HTML",
        "CSS",
        "JavaScript"
    ]

    for skill in skill_list:
        if skill.lower() in text.lower():
            skills.append(skill)

    questions = []

    for skill in skills:

        if skill == "Python":
            questions.extend([
                "What is Python?",
                "Explain Python decorators.",
                "Difference between List and Tuple?"
            ])

        elif skill == "Java":
            questions.extend([
                "What is JVM?",
                "Explain OOP concepts in Java.",
                "Difference between JDK and JRE?"
            ])

        elif skill == "Machine Learning":
            questions.extend([
                "What is Machine Learning?",
                "What is Overfitting?",
                "Explain Gradient Boosting."
            ])

        elif skill == "SQL":
            questions.extend([
                "Explain SQL Joins.",
                "Difference between DELETE and TRUNCATE?"
            ])

    return {
        "skills": skills,
        "questions": questions
    }
@app.post("/evaluate-answer")
def evaluate_answer(
    data: AnswerRequest,
    db: Session = Depends(get_db)
):

    answer_length = len(data.answer.split())

    technical_score = 5
    communication_score = 5

    if answer_length > 10:
        technical_score += 2
        communication_score += 2

    if answer_length > 20:
        technical_score += 2
        communication_score += 1

    overall_score = round(
        (technical_score + communication_score) / 2,
        1
    )

    feedback = "Good answer."

    if overall_score < 7:
        feedback = "Try giving a more detailed explanation."

    elif overall_score >= 8:
        feedback = "Excellent answer. Add practical examples for even better impact."

    history = InterviewHistory(
        question=data.question,
        answer=data.answer,
        technical_score=technical_score,
        communication_score=communication_score,
        overall_score=overall_score
    )

    db.add(history)
    db.commit()

    return {
        "question": data.question,
        "technical_score": technical_score,
        "communication_score": communication_score,
        "overall_score": overall_score,
        "feedback": feedback
    }
@app.get("/interview-history")
def get_history(db: Session = Depends(get_db)):

    history = db.query(InterviewHistory).all()

    result = []

    for item in history:
        result.append({
            "question": item.question,
            "answer": item.answer,
            "technical_score": item.technical_score,
            "communication_score": item.communication_score,
            "overall_score": item.overall_score
        })

    return result

def dashboard(db: Session = Depends(get_db)):


    total_score = 0

    for item in interviews:
        total_score += item.overall_score

    average_score = round(
        total_score / total_interviews,
        1
    )

    strong_area = "Technical Skills"
    weak_area = "Communication"

    return {
        "total_interviews": total_interviews,
        "average_score": average_score,
        "strong_area": strong_area,
        "weak_area": weak_area
    }
@app.post("/resume-score")
async def resume_score(file: UploadFile = File(...)):

    text = ""

    if file.filename.endswith(".pdf"):

        pdf_data = await file.read()
        pdf = fitz.open(stream=pdf_data, filetype="pdf")

        for page in pdf:
            text += page.get_text()

    elif file.filename.endswith(".docx"):

        doc = Document(file.file)

        for para in doc.paragraphs:
            text += para.text + "\n"

    score = 50
    suggestions = []

    if "github" in text.lower():
        score += 10
    else:
        suggestions.append("Add GitHub profile")

    if "project" in text.lower():
        score += 15
    else:
        suggestions.append("Add Projects")

    if "certificate" in text.lower():
        score += 10
    else:
        suggestions.append("Add Certifications")

    if "python" in text.lower():
        score += 5

    if "java" in text.lower():
        score += 5

    if score > 100:
        score = 100

    return {
        "resume_score": score,
        "suggestions": suggestions
    }
@app.post("/ai-generate-questions")
def ai_generate_questions(data: QuestionRequest):

    try:

        prompt = f"""
        Generate 5 technical interview questions
        for these skills:

        {', '.join(data.skills)}

        Return only the questions.
        """

        print("API KEY =", os.getenv("GEMINI_API_KEY"))

        response = model.generate_content(prompt)

        return {
            "questions": response.text
        }

    except Exception as e:

        print("ERROR =", str(e))

        return {
            "error": str(e)
        }


@app.get("/list-models")
def list_models():

    models = []

    for m in genai.list_models():
        models.append(m.name)

    return {
        "models": models
    }
@app.post("/ai-evaluate-answer")
def ai_evaluate_answer(data: AIAnswerRequest):

    prompt = f"""
    You are an interview evaluator.

    Question:
    {data.question}

    Candidate Answer:
    {data.answer}

    Evaluate the answer and provide:

    1. Technical Score (out of 10)
    2. Communication Score (out of 10)
    3. Overall Score (out of 10)
    4. Feedback

    Return in simple text format.
    """

    response = model.generate_content(prompt)

    return {
        "evaluation": response.text
    }
@app.post("/voice-interview")
async def voice_interview(file: UploadFile = File(...)):

    try:

        temp_file = tempfile.NamedTemporaryFile(delete=False)

        contents = await file.read()

        temp_file.write(contents)
        temp_file.close()

        recognizer = sr.Recognizer()

        with sr.AudioFile(temp_file.name) as source:
            audio = recognizer.record(source)

        text = recognizer.recognize_google(audio)

        prompt = f"""
        You are an interview evaluator.

        Candidate Answer:
        {text}

        Evaluate the answer and provide:

        1. Technical Score (out of 10)
        2. Communication Score (out of 10)
        3. Overall Score (out of 10)
        4. Feedback
        """

        response = model.generate_content(prompt)

        return {
            "transcribed_text": text,
            "evaluation": response.text
        }

    except Exception as e:

        return {
            "error": str(e)
        }
@app.get("/hr-questions")
def hr_questions():

    questions = [
        "Tell me about yourself.",
        "Why should we hire you?",
        "What are your strengths?",
        "What are your weaknesses?",
        "Where do you see yourself in 5 years?",
        "Why do you want to join our company?",
        "Describe a challenging situation you faced.",
        "How do you handle pressure?",
        "What motivates you?",
        "Tell me about a project you are proud of."
    ]

    return {
        "questions": questions
    }
@app.post("/resume-analysis")
async def resume_analysis(file: UploadFile = File(...)):

    pdf = fitz.open(stream=await file.read(), filetype="pdf")

    text = ""

    for page in pdf:
        text += page.get_text()

    prompt = f"""
You are an expert HR recruiter.

Analyze the following resume.

Resume:

{text}

Provide:
1. ATS Score (out of 100)
2. Resume Score (out of 10)
3. Technical Skills
4. Soft Skills
5. Strengths
6. Weaknesses
7. Missing Skills
8. Suggested Improvements
9. Suitable Job Roles
10. Five Interview Questions
"""

    response = model.generate_content(prompt)

    return {
        "analysis": response.text
    }
from pydantic import BaseModel

class InterviewAnswer(BaseModel):
    answer: str


@app.post("/evaluate-answer")
async def evaluate_answer(data: InterviewAnswer):

    prompt = f"""
You are an HR interviewer.

Evaluate this interview answer:

{data.answer}

Provide:

1. Communication Score (/10)

2. Confidence Score (/10)

3. Technical Score (/10)

4. Strengths

5. Weaknesses

6. Suggestions

7. Overall Feedback
"""

    response = model.generate_content(prompt)

    return {
        "feedback": response.text
    }
from pydantic import BaseModel

class CodeRequest(BaseModel):
    question: str
    code: str


class ResumeRequest(BaseModel):
    name: str
    email: str
    skills: str
    projects: str


@app.post("/evaluate-code")
async def evaluate_code(data: CodeRequest):

    prompt = f"""
You are an expert technical interviewer.

Coding Question:
{data.question}

Candidate's Code:
{data.code}

Evaluate the code and provide:

1. Code Score (out of 10)

2. Correctness

3. Time Complexity

4. Space Complexity

5. Strengths

6. Mistakes

7. Suggested Improvements

8. Optimized Solution
"""

    response = model.generate_content(prompt)

    return {
        "feedback": response.text
    }
@app.post("/generate-resume")
async def generate_resume(data: ResumeRequest):

    prompt = f"""
Create a professional ATS-friendly resume.

Name:
{data.name}

Email:
{data.email}

Skills:
{data.skills}

Projects:
{data.projects}

Write a professional resume with:

1. Career Objective
2. Professional Summary
3. Technical Skills
4. Projects
5. Strengths

Return only the resume.
"""

    try:
        response = model.generate_content(prompt)

        return {
            "resume": response.text
        }

    except Exception:
        return {
            "resume": """
⚠ AI Resume Generation is temporarily unavailable.

Reason:
Gemini API quota has been exceeded.

Please try again later or use a new Gemini API key.
"""
        }